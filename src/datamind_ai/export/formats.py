from __future__ import annotations

import io

import pandas as pd
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

from datamind_ai.session import DatasetInfo


class ReportPDF(FPDF):
    def header(self) -> None:
        self.set_font("Helvetica", "B", 14)
        self.cell(0, 10, "DataMind AI - Relatorio", new_x="LMARGIN", new_y="NEXT", align="C")
        self.ln(4)

    def footer(self) -> None:
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Pagina {self.page_no()}", align="C")


def _latin1_safe(text: str) -> str:
    replacements = {
        "á": "a", "à": "a", "ã": "a", "â": "a",
        "é": "e", "ê": "e",
        "í": "i",
        "ó": "o", "õ": "o", "ô": "o",
        "ú": "u", "ç": "c",
        "Á": "A", "À": "A", "Ã": "A", "Â": "A",
        "É": "E", "Ê": "E",
        "Í": "I",
        "Ó": "O", "Õ": "O", "Ô": "O",
        "Ú": "U", "Ç": "C",
        "—": "-", "⚠": "!", "✅": "OK",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def export_to_excel(info: DatasetInfo) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        overview = info.overview or {}
        summary_data = {
            "Metrica": ["Linhas", "Colunas", "Duplicados", "% Duplicados"],
            "Valor": [
                overview.get("row_count", 0),
                overview.get("column_count", 0),
                overview.get("duplicate_rows", 0),
                overview.get("duplicate_pct", 0),
            ],
        }
        pd.DataFrame(summary_data).to_excel(writer, sheet_name="Resumo", index=False)

        if overview.get("columns"):
            pd.DataFrame(overview["columns"]).to_excel(
                writer, sheet_name="Colunas", index=False
            )

        if info.quality_report and info.quality_report.get("issues"):
            pd.DataFrame(info.quality_report["issues"]).to_excel(
                writer, sheet_name="Qualidade", index=False
            )

        if info.data_dictionary:
            pd.DataFrame(info.data_dictionary).to_excel(
                writer, sheet_name="Dicionario", index=False
            )

        if info.statistics:
            num_df, cat_df = pd.DataFrame(info.statistics.get("numeric", [])), pd.DataFrame(
                info.statistics.get("categorical", [])
            )
            if not num_df.empty:
                num_df.to_excel(writer, sheet_name="Stats Numericas", index=False)
            if not cat_df.empty:
                cat_df.to_excel(writer, sheet_name="Stats Categoricas", index=False)

        if info.business_rules:
            pd.DataFrame(info.business_rules).to_excel(
                writer, sheet_name="Regras Negocio", index=False
            )

        if info.kpi_suggestions:
            pd.DataFrame(info.kpi_suggestions).to_excel(
                writer, sheet_name="KPIs", index=False
            )

        info.df.head(1000).to_excel(writer, sheet_name="Dados Amostra", index=False)

    buffer.seek(0)
    return buffer.getvalue()


def _pdf_text(pdf: ReportPDF, text: str, h: float = 5) -> None:
    if pdf.get_y() > 270:
        pdf.add_page()
    w = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(w, h, _latin1_safe(text))


def export_to_pdf(info: DatasetInfo) -> bytes:
    overview = info.overview or {}
    pdf = ReportPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    _pdf_text(pdf, f"Relatorio DataMind AI - {info.name}", 8)
    pdf.set_font("Helvetica", size=10)
    _pdf_text(pdf, f"Ficheiro: {info.source_filename}", 6)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    _pdf_text(pdf, "Resumo Executivo", 8)
    pdf.set_font("Helvetica", size=10)
    summary = (
        f"Linhas: {overview.get('row_count', '?')} | "
        f"Colunas: {overview.get('column_count', '?')} | "
        f"Duplicados: {overview.get('duplicate_rows', 0)}"
    )
    _pdf_text(pdf, summary, 6)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    _pdf_text(pdf, "Qualidade de Dados", 8)
    pdf.set_font("Helvetica", size=10)
    if info.quality_report and info.quality_report.get("issues"):
        for issue in info.quality_report["issues"][:20]:
            _pdf_text(pdf, f"- {issue['category']} ({issue['column']}): {issue['pct']}%")
    else:
        _pdf_text(pdf, "Nenhum problema detetado.")
    pdf.ln(4)

    if info.business_rules:
        pdf.set_font("Helvetica", "B", 12)
        _pdf_text(pdf, "Regras de Negocio (hipoteses)", 8)
        pdf.set_font("Helvetica", size=10)
        for rule in info.business_rules[:15]:
            _pdf_text(pdf, f"- {rule.get('rule', '')} ({rule.get('evidence_pct', 0)}%)")
        pdf.ln(4)

    if info.kpi_suggestions:
        pdf.set_font("Helvetica", "B", 12)
        _pdf_text(pdf, "KPIs Sugeridos", 8)
        pdf.set_font("Helvetica", size=10)
        for kpi in info.kpi_suggestions[:15]:
            _pdf_text(pdf, f"- {kpi.get('name', '')}: {kpi.get('formula', '')}")

    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    _pdf_text(pdf, "Relatorio gerado pelo DataMind AI. Validar interpretacoes de IA.", 4)

    return bytes(pdf.output())


def export_to_word(info: DatasetInfo) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    overview = info.overview or {}
    doc.add_heading(f"Relatório DataMind AI — {info.name}", level=0)
    doc.add_paragraph(f"Ficheiro: {info.source_filename}")

    doc.add_heading("Resumo Executivo", level=1)
    doc.add_paragraph(
        f"O dataset {info.name} contém {overview.get('row_count', '?')} linhas e "
        f"{overview.get('column_count', '?')} colunas. "
        f"Duplicados: {overview.get('duplicate_rows', 0)} ({overview.get('duplicate_pct', 0)}%)."
    )

    doc.add_heading("Colunas", level=1)
    if overview.get("columns"):
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Coluna", "Tipo", "Valores em falta", "% em falta"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for col in overview["columns"]:
            row = table.add_row().cells
            row[0].text = str(col.get("column", ""))
            row[1].text = str(col.get("dtype", ""))
            row[2].text = str(col.get("null_count", ""))
            row[3].text = f"{col.get('null_pct', '')}%"

    doc.add_heading("Qualidade de Dados", level=1)
    if info.quality_report and info.quality_report.get("issues"):
        for issue in info.quality_report["issues"]:
            doc.add_paragraph(
                f"{issue['category']} — {issue['column']}: "
                f"{issue['count']} ocorrências ({issue['pct']}%). "
                f"Recomendação: {issue['recommendation']}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("Nenhum problema detetado.")

    doc.add_heading("Dicionário de Dados", level=1)
    if info.data_dictionary:
        for entry in info.data_dictionary:
            doc.add_paragraph(
                f"{entry.get('column', '')} ({entry.get('dtype', '')}): "
                f"{entry.get('description', '')} — "
                f"Sugestão: {entry.get('business_meaning', '')}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("Dicionário não gerado.")

    if info.business_rules:
        doc.add_heading("Regras de Negócio (hipóteses)", level=1)
        for rule in info.business_rules:
            doc.add_paragraph(
                f"{rule.get('rule', '')} — {rule.get('evidence_pct', 0)}% suporte",
                style="List Bullet",
            )

    if info.kpi_suggestions:
        doc.add_heading("KPIs Sugeridos", level=1)
        for kpi in info.kpi_suggestions:
            avail = "Sim" if kpi.get("columns_available") else "Não"
            doc.add_paragraph(
                f"{kpi.get('name', '')}: {kpi.get('formula', '')} "
                f"(Colunas disponíveis: {avail})",
                style="List Bullet",
            )

    doc.add_paragraph(
        "Relatório gerado automaticamente pelo DataMind AI. "
        "Interpretações de IA devem ser validadas."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
